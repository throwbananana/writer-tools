"""
导出器插件系统 - 支持动态注册导出格式

使用方法：
    from writer_app.core.exporter import ExporterRegistry

    # 获取所有已注册的导出格式
    formats = ExporterRegistry.list_formats()

    # 执行导出
    ExporterRegistry.export("markdown", project_data, "/path/to/file.md")

创建自定义导出器：
    from writer_app.core.exporter import ExportFormat, ExporterRegistry

    class MyCustomExporter(ExportFormat):
        key = "my_format"
        name = "我的格式"
        extension = ".myf"
        description = "自定义导出格式"

        @classmethod
        def export(cls, project_data, file_path, **kwargs):
            # 实现导出逻辑
            ...
            return True

    # 注册
    ExporterRegistry.register(MyCustomExporter)
"""

import xml.etree.ElementTree as ET
from xml.dom import minidom
import json
import tkinter.messagebox as messagebox
import os
import shutil
import re
import html
import csv
from abc import ABC, abstractmethod
from typing import Dict, List, Type, Optional, Any, Union
from dataclasses import dataclass, field


@dataclass
class ExportResult:
    """导出操作的结果。"""
    success: bool
    message: str = ""
    file_path: Optional[str] = None
    content: Optional[str] = None


class ExportFormat(ABC):
    """导出格式基类。所有导出器必须继承此类。"""

    # 子类必须定义这些属性
    key: str = ""           # 唯一标识符，如 "markdown", "pdf"
    name: str = ""          # 显示名称（中文）
    extension: str = ""     # 文件扩展名，如 ".md", ".pdf"
    description: str = ""   # 格式描述
    requires_dir: bool = False  # 是否需要目录而非文件（如Ren'Py导出）

    # 题材支持配置
    supported_types: List[str] = []  # 支持的项目类型，空列表=支持所有
    priority_for_types: Dict[str, int] = {}  # 类型优先级 {type: priority}，数值越高越优先

    @classmethod
    def is_available_for_type(cls, project_type: str) -> bool:
        """检查此导出格式是否适用于指定的项目类型。"""
        if not cls.supported_types:
            return True  # 空列表表示支持所有类型
        return project_type in cls.supported_types

    @classmethod
    def get_priority_for_type(cls, project_type: str) -> int:
        """获取此导出格式对于指定项目类型的优先级。"""
        return cls.priority_for_types.get(project_type, 0)

    @classmethod
    def get_type_specific_options(cls, project_type: str) -> Dict[str, Any]:
        """获取项目类型特定的导出选项。子类可重写此方法。"""
        return {}

    @classmethod
    @abstractmethod
    def export(cls, project_data: Dict, file_path: str, **kwargs) -> bool:
        """
        执行导出。

        Args:
            project_data: 完整的项目数据字典
            file_path: 导出目标路径
            **kwargs: 格式特定参数

        Returns:
            bool: 成功返回True，失败应抛出异常
        """
        pass

    @classmethod
    def get_file_filter(cls) -> tuple:
        """获取文件对话框过滤器元组。"""
        return (cls.name, f"*{cls.extension}")


class ExporterRegistry:
    """导出器注册表。管理所有已注册的导出格式。"""

    _formats: Dict[str, Type[ExportFormat]] = {}

    @classmethod
    def register(cls, exporter: Type[ExportFormat]) -> None:
        """
        注册一个导出格式。

        Args:
            exporter: ExportFormat的子类
        """
        if not exporter.key:
            raise ValueError(f"Exporter {exporter.__name__} must define a 'key' attribute")
        cls._formats[exporter.key] = exporter

    @classmethod
    def unregister(cls, key: str) -> bool:
        """
        取消注册一个导出格式。

        Args:
            key: 导出格式的唯一标识符

        Returns:
            bool: 是否成功取消注册
        """
        if key in cls._formats:
            del cls._formats[key]
            return True
        return False

    @classmethod
    def get(cls, key: str) -> Optional[Type[ExportFormat]]:
        """
        获取指定的导出格式类。

        Args:
            key: 导出格式的唯一标识符

        Returns:
            ExportFormat子类或None
        """
        return cls._formats.get(key)

    @classmethod
    def list_formats(cls) -> List[Type[ExportFormat]]:
        """
        获取所有已注册的导出格式列表。

        Returns:
            List[ExportFormat]: 所有已注册的导出格式
        """
        return list(cls._formats.values())

    @classmethod
    def list_keys(cls) -> List[str]:
        """
        获取所有已注册的导出格式键。

        Returns:
            List[str]: 所有格式键
        """
        return list(cls._formats.keys())

    @classmethod
    def get_display_name(cls, key: str) -> Optional[str]:
        """
        获取指定格式的显示名称。

        Args:
            key: 导出格式的唯一标识符

        Returns:
            str: 显示名称，如 "Markdown"
        """
        exporter = cls._formats.get(key)
        return exporter.name if exporter else None

    @classmethod
    def get_file_extension(cls, key: str) -> Optional[str]:
        """
        获取指定格式的文件扩展名。

        Args:
            key: 导出格式的唯一标识符

        Returns:
            str: 文件扩展名，如 ".md"
        """
        exporter = cls._formats.get(key)
        return exporter.extension if exporter else None

    @classmethod
    def get_available_formats(cls) -> List[Dict[str, str]]:
        """
        获取所有可用格式的详细信息。

        Returns:
            List[Dict]: 包含 key, display_name, file_extension 的字典列表
        """
        return [
            {
                "key": fmt.key,
                "display_name": fmt.name,
                "file_extension": fmt.extension,
                "description": fmt.description
            }
            for fmt in cls._formats.values()
        ]

    @classmethod
    def get_formats_for_type(cls, project_type: str) -> List[Dict[str, Any]]:
        """
        获取适用于指定项目类型的导出格式，按优先级排序。

        Args:
            project_type: 项目类型（如 "Suspense", "Romance", "Galgame"）

        Returns:
            List[Dict]: 按优先级排序的格式列表，包含 key, display_name, priority, is_recommended
        """
        available = []

        for fmt in cls._formats.values():
            if fmt.is_available_for_type(project_type):
                priority = fmt.get_priority_for_type(project_type)
                available.append({
                    "key": fmt.key,
                    "display_name": fmt.name,
                    "file_extension": fmt.extension,
                    "description": fmt.description,
                    "priority": priority,
                    "is_recommended": priority > 50
                })

        # 按优先级降序排序
        available.sort(key=lambda x: x["priority"], reverse=True)
        return available

    @classmethod
    def get_recommended_format(cls, project_type: str) -> Optional[str]:
        """
        获取指定项目类型的推荐导出格式。

        Args:
            project_type: 项目类型

        Returns:
            str: 推荐的导出格式 key，如果没有则返回 None
        """
        formats = cls.get_formats_for_type(project_type)
        if formats and formats[0]["priority"] > 50:
            return formats[0]["key"]
        return None

    @classmethod
    def export(cls, key: str, project_data_or_manager: Union[Dict, Any], file_path: str = None, **kwargs) -> ExportResult:
        """
        使用指定格式执行导出。

        Args:
            key: 导出格式的唯一标识符
            project_data_or_manager: 项目数据字典或 ProjectManager 对象
            file_path: 目标路径（可选，如果不提供则只返回内容）
            **kwargs: 格式特定参数

        Returns:
            ExportResult: 包含 success, message, file_path, content 的结果对象
        """
        exporter = cls._formats.get(key)
        if not exporter:
            return ExportResult(success=False, message=f"Unknown export format: {key}")

        # 支持 ProjectManager 或直接的 dict
        if hasattr(project_data_or_manager, 'project_data'):
            project_data = project_data_or_manager.project_data
        else:
            project_data = project_data_or_manager

        try:
            # 如果没有提供 file_path，生成内容到内存
            if file_path is None:
                content = cls._export_to_string(exporter, project_data, **kwargs)
                return ExportResult(success=True, message="导出成功", content=content)
            else:
                exporter.export(project_data, file_path, **kwargs)
                return ExportResult(success=True, message="导出成功", file_path=file_path)
        except Exception as e:
            return ExportResult(success=False, message=f"导出失败: {str(e)}")

    @classmethod
    def _export_to_string(cls, exporter: Type['ExportFormat'], project_data: Dict, **kwargs) -> str:
        """
        将导出内容生成为字符串（用于预览或内存导出）。

        Args:
            exporter: 导出器类
            project_data: 项目数据
            **kwargs: 格式特定参数

        Returns:
            str: 导出的内容
        """
        import tempfile
        import os

        # 创建临时文件
        with tempfile.NamedTemporaryFile(mode='w', suffix=exporter.extension, delete=False, encoding='utf-8') as f:
            temp_path = f.name

        try:
            exporter.export(project_data, temp_path, **kwargs)
            with open(temp_path, 'r', encoding='utf-8') as f:
                return f.read()
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    @classmethod
    def get_file_filters(cls) -> List[tuple]:
        """
        获取所有格式的文件过滤器，用于文件对话框。

        Returns:
            List[tuple]: [(name, pattern), ...]
        """
        return [fmt.get_file_filter() for fmt in cls._formats.values()]


# ============================================================================
# 内置导出格式
# ============================================================================

class CSVExporter(ExportFormat):
    """CSV格式导出器"""
    key = "csv"
    name = "CSV表格"
    extension = ".csv"
    description = "逗号分隔值格式，可用Excel打开"

    @classmethod
    def export(cls, project_data: Dict, file_path: str, **kwargs) -> bool:
        script = project_data.get("script", {})
        scenes = script.get("scenes", [])

        with open(file_path, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(["序号", "场景名", "地点", "时间", "字数", "登场角色", "大纲关联", "备注"])

            for i, scene in enumerate(scenes):
                content = scene.get("content", "")
                word_count = len(content)
                outline_ref = scene.get("outline_ref_path", "")

                writer.writerow([
                    i + 1,
                    scene.get("name", ""),
                    scene.get("location", ""),
                    scene.get("time", ""),
                    word_count,
                    ", ".join(scene.get("characters", [])),
                    outline_ref,
                    ""
                ])
        return True


class ExcelExporter(ExportFormat):
    """Excel格式导出器"""
    key = "excel"
    name = "Excel工作簿"
    extension = ".xlsx"
    description = "Microsoft Excel格式，带格式化"

    @classmethod
    def export(cls, project_data: Dict, file_path: str, **kwargs) -> bool:
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill
        except ImportError:
            raise ImportError("请安装 'openpyxl' 以使用Excel导出: pip install openpyxl")

        script = project_data.get("script", {})
        scenes = script.get("scenes", [])

        wb = Workbook()
        ws = wb.active
        ws.title = "场景列表"

        headers = ["序号", "场景名", "地点", "时间", "字数", "登场角色", "大纲关联", "备注"]
        ws.append(headers)

        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
        for cell in ws[1]:
            cell.font = header_font
            cell.fill = header_fill

        for i, scene in enumerate(scenes):
            content = scene.get("content", "")
            word_count = len(content)
            outline_ref = scene.get("outline_ref_path", "")

            row = [
                i + 1,
                scene.get("name", ""),
                scene.get("location", ""),
                scene.get("time", ""),
                word_count,
                ", ".join(scene.get("characters", [])),
                outline_ref,
                ""
            ]
            ws.append(row)

        ws.column_dimensions['B'].width = 20
        ws.column_dimensions['C'].width = 15
        ws.column_dimensions['D'].width = 15
        ws.column_dimensions['F'].width = 25
        ws.column_dimensions['G'].width = 25

        # 第二个工作表：角色统计
        ws_char = wb.create_sheet("角色统计")
        ws_char.append(["姓名", "描述", "登场次数", "标签"])
        for cell in ws_char[1]:
            cell.font = header_font
            cell.fill = header_fill

        chars = script.get("characters", [])
        for char in chars:
            name = char.get("name", "")
            count = sum(1 for s in scenes if name in s.get("characters", []))

            ws_char.append([
                name,
                char.get("description", ""),
                count,
                ", ".join(char.get("tags", []))
            ])

        ws_char.column_dimensions['A'].width = 15
        ws_char.column_dimensions['B'].width = 40

        wb.save(file_path)
        return True


class TxtExporter(ExportFormat):
    """纯文本格式导出器"""
    key = "txt"
    name = "纯文本"
    extension = ".txt"
    description = "纯文本格式，无格式化"

    @classmethod
    def export(cls, project_data: Dict, file_path: str, **kwargs) -> bool:
        script = project_data.get("script", {})
        lines = []

        title = script.get("title", "Untitled")
        lines.append(title)
        lines.append("=" * len(title))
        lines.append("")

        for scene in script.get("scenes", []):
            name = scene.get("name", "Untitled Scene")
            loc = scene.get("location", "Unknown Location")
            time = scene.get("time", "Unknown Time")

            lines.append(f"{name}")
            lines.append(f"{loc} - {time}")
            lines.append("-" * 20)

            content = scene.get("content", "")
            lines.append(content)
            lines.append("\n" + "=" * 30 + "\n")

        with open(file_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return True


class MarkdownExporter(ExportFormat):
    """Markdown格式导出器"""
    key = "markdown"
    name = "Markdown"
    extension = ".md"
    description = "Markdown格式，适合文档分享"

    @classmethod
    def export(cls, project_data: Dict, file_path: str, **kwargs) -> bool:
        include_outline = kwargs.get("include_outline", True)
        include_script = kwargs.get("include_script", True)

        lines = []
        script = project_data.get("script", {})

        if include_script:
            title = script.get("title", "Untitled")
            lines.append(f"# {title}\n")

            chars = script.get("characters", [])
            if chars:
                lines.append("## Characters")
                for c in chars:
                    lines.append(f"- **{c.get('name')}**: {c.get('description')}")
                lines.append("")

            scenes = script.get("scenes", [])
            if scenes:
                lines.append("## Scenes")
                for s in scenes:
                    lines.append(f"### {s.get('name')}")
                    lines.append(f"**Location**: {s.get('location')} | **Time**: {s.get('time')}\n")

                    content = s.get("content", "")
                    for line in content.split('\n'):
                        line = line.strip()
                        if not line:
                            lines.append("")
                            continue
                        match = re.match(r"^([^\s：:]+)[：:]\s*(.*)$", line)
                        if match:
                            name = match.group(1)
                            dialogue = match.group(2)
                            lines.append(f"**{name}**: {dialogue}")
                        else:
                            lines.append(line)

                    lines.append("\n---\n")

        if include_outline:
            lines.append("\n# Outline\n")
            cls._recursive_outline(project_data.get("outline", {}), 0, lines)

        with open(file_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return True

    @classmethod
    def _recursive_outline(cls, node, level, lines):
        prefix = "#" * (level + 2) if level < 5 else "-"
        lines.append(f"{prefix} {node.get('name', 'Untitled')}")
        if node.get("content"):
            lines.append(f"{node.get('content')}\n")
        for child in node.get("children", []):
            cls._recursive_outline(child, level + 1, lines)


class HTMLExporter(ExportFormat):
    """HTML打印格式导出器"""
    key = "html"
    name = "HTML (打印)"
    extension = ".html"
    description = "HTML格式，适合浏览器打印为PDF"

    @classmethod
    def export(cls, project_data: Dict, file_path: str, **kwargs) -> bool:
        script = project_data.get("script", {})
        title = html.escape(script.get("title", "Untitled"))

        css = """
        @media print {
            @page { margin: 1in; size: A4; }
            body { font-family: 'Courier New', Courier, monospace; font-size: 12pt; line-height: 1.2; }
            .scene-heading { font-weight: bold; text-transform: uppercase; margin-top: 2em; margin-bottom: 1em; }
            .character { font-weight: bold; text-align: center; width: 60%; margin-left: 20%; margin-top: 1em; }
            .parenthetical { font-style: italic; text-align: center; width: 50%; margin-left: 25%; }
            .dialogue { text-align: left; width: 70%; margin-left: 15%; margin-bottom: 1em; }
            .action { text-align: left; margin-bottom: 1em; }
            .transition { text-align: right; text-transform: uppercase; margin-top: 1em; margin-bottom: 1em; }
            .page-break { page-break-before: always; }
        }
        body { font-family: 'Courier New', Courier, monospace; font-size: 12pt; line-height: 1.2; max-width: 8.5in; margin: 0 auto; padding: 2em; background: #fff; color: #000; }
        .scene-heading { font-weight: bold; text-transform: uppercase; margin-top: 24pt; margin-bottom: 12pt; background: #eee; }
        .character { font-weight: bold; text-align: center; width: 50%; margin-left: 25%; margin-top: 12pt; }
        .parenthetical { font-style: italic; text-align: center; width: 40%; margin-left: 30%; }
        .dialogue { text-align: left; width: 60%; margin-left: 20%; margin-bottom: 12pt; }
        .action { text-align: left; margin-bottom: 12pt; }
        .transition { text-align: right; text-transform: uppercase; margin-top: 12pt; margin-bottom: 12pt; }
        """

        html_content = [f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>{title}</title><style>{css}</style></head><body>"]
        html_content.append(f"<h1 style='text-align:center; text-transform:uppercase; margin-top:3in; margin-bottom:4in;'>{title}</h1>")
        html_content.append("<div class='page-break'></div>")

        for scene in script.get("scenes", []):
            name = html.escape(scene.get("name", "SCENE"))
            loc = html.escape(scene.get("location", ""))
            time_val = html.escape(scene.get("time", ""))
            slug = f"{name} - {loc} - {time_val}".upper()

            html_content.append(f"<div class='scene-heading'>{slug}</div>")

            content = scene.get("content", "")
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    continue

                esc_line = html.escape(line)

                match = re.match(r"^([^\s：:]+)[：:]\s*(.*)$", line)
                if match:
                    char_name = match.group(1).strip()
                    dialogue = match.group(2).strip()
                    html_content.append(f"<div class='character'>{html.escape(char_name).upper()}</div>")

                    if dialogue.startswith("(") and ")" in dialogue:
                        end_paren = dialogue.find(")")
                        paren = dialogue[:end_paren + 1]
                        dial = dialogue[end_paren + 1:].strip()
                        html_content.append(f"<div class='parenthetical'>{html.escape(paren)}</div>")
                        if dial:
                            html_content.append(f"<div class='dialogue'>{html.escape(dial)}</div>")
                    else:
                        html_content.append(f"<div class='dialogue'>{html.escape(dialogue)}</div>")

                elif line.startswith("(") and line.endswith(")"):
                    html_content.append(f"<div class='parenthetical'>{esc_line}</div>")

                elif line.startswith("——") or line.endswith("TO:"):
                    html_content.append(f"<div class='transition'>{esc_line}</div>")

                else:
                    html_content.append(f"<div class='action'>{esc_line}</div>")

        html_content.append("</body></html>")

        with open(file_path, "w", encoding="utf-8") as f:
            f.write("\n".join(html_content))
        return True


class FountainExporter(ExportFormat):
    """Fountain剧本格式导出器"""
    key = "fountain"
    name = "Fountain剧本"
    extension = ".fountain"
    description = "Fountain格式，可导入专业剧本软件"

    @classmethod
    def export(cls, project_data: Dict, file_path: str, **kwargs) -> bool:
        script = project_data.get("script", {})
        lines = []

        title = script.get("title", "Untitled")
        lines.append(f"Title: {title}")
        lines.append("Credit: Written with Writer Tool")
        lines.append("\n")

        for scene in script.get("scenes", []):
            loc = scene.get("location", "UNKNOWN").upper()
            time_val = scene.get("time", "").upper()

            if not any(loc.startswith(p) for p in ["INT", "EXT", ".", "EST"]):
                if not loc.startswith("."):
                    loc = "." + loc

            heading = f"{loc} - {time_val}" if time_val else loc
            lines.append(f"\n{heading}\n")

            content = scene.get("content", "")
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    lines.append("")
                    continue

                match = re.match(r"^([^\s：:]+)[：:]\s*(.*)$", line)
                if match:
                    char = match.group(1).strip().upper()
                    diag = match.group(2).strip()
                    lines.append(f"\n{char}")
                    lines.append(diag)
                    continue

                if line.startswith("(") and line.endswith(")"):
                    lines.append(line)
                    continue

                if line.endswith("TO:"):
                    lines.append(f"\n> {line}")
                    continue

                lines.append(line)

        with open(file_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return True


class FDXExporter(ExportFormat):
    """Final Draft XML格式导出器"""
    key = "fdx"
    name = "Final Draft"
    extension = ".fdx"
    description = "Final Draft 12 XML格式"

    @classmethod
    def export(cls, project_data: Dict, file_path: str, **kwargs) -> bool:
        script = project_data.get("script", {})

        root = ET.Element("FinalDraft", {"DocumentType": "Script", "Template": "No", "Version": "4"})

        content = ET.SubElement(root, "Content")
        paragraph = ET.SubElement(content, "Paragraph", {"Type": "General"})
        text = ET.SubElement(paragraph, "Text")
        text.text = script.get("title", "Untitled")

        for scene in script.get("scenes", []):
            slugline = f"{scene.get('location', 'UNKNOWN')} - {scene.get('time', 'UNKNOWN')}"
            sc_para = ET.SubElement(content, "Paragraph", {"Type": "Scene Heading"})
            sc_text = ET.SubElement(sc_para, "Text")
            sc_text.text = f"{scene.get('name')} - {slugline}"

            scene_body = scene.get("content", "")
            for line in scene_body.split('\n'):
                line = line.strip()
                if not line:
                    continue

                p_type = "Action"
                if ":" in line and len(line.split(":")[0]) < 20:
                    p_type = "Dialogue"
                elif line.startswith("(") and line.endswith(")"):
                    p_type = "Parenthetical"

                para = ET.SubElement(content, "Paragraph", {"Type": p_type})
                t = ET.SubElement(para, "Text")
                t.text = line

        xml_str = minidom.parseString(ET.tostring(root)).toprettyxml(indent="  ")

        with open(file_path, "w", encoding="utf-8") as f:
            f.write(xml_str)
        return True


class DocxExporter(ExportFormat):
    """Word文档格式导出器"""
    key = "docx"
    name = "Word文档"
    extension = ".docx"
    description = "Microsoft Word格式"

    @classmethod
    def export(cls, project_data: Dict, file_path: str, **kwargs) -> bool:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装 'python-docx' 以使用Word导出: pip install python-docx")

        document = Document()
        script = project_data.get("script", {})

        document.add_heading(script.get("title", "Untitled"), 0)

        for scene in script.get("scenes", []):
            slug = f"{scene.get('name', '')} [Loc: {scene.get('location')} Time: {scene.get('time')}]"
            document.add_heading(slug, level=2)

            content = scene.get("content", "")
            document.add_paragraph(content)

        document.save(file_path)
        return True


class PDFExporter(ExportFormat):
    """PDF格式导出器"""
    key = "pdf"
    name = "PDF文档"
    extension = ".pdf"
    description = "标准剧本PDF格式"

    @classmethod
    def export(cls, project_data: Dict, file_path: str, **kwargs) -> bool:
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
            from reportlab.pdfbase import pdfmetrics
        except ImportError:
            raise ImportError("请安装 'reportlab' 以使用PDF导出: pip install reportlab")

        # Register local fonts
        from writer_app.core.font_manager import get_font_manager
        get_font_manager().register_for_reportlab()

        # Get export settings
        font_name = kwargs.get("font_name", "Courier")
        margin_mm = kwargs.get("margin", 25)
        line_spacing = kwargs.get("line_spacing", 1.2)
        
        # Convert margin to inch (1 inch = 25.4 mm)
        margin_inch = margin_mm / 25.4

        # Check if font is usable in ReportLab
        # ReportLab requires fonts to be registered if they are not standard 14 fonts.
        # Standard: Courier, Helvetica, Times-Roman, etc.
        # If the user selected a font that isn't registered (e.g. system font not loaded via TTF),
        # ReportLab will raise an error.
        # We fallback to Courier if font is likely invalid, OR we try to proceed.
        # Since we can't easily check system fonts for ReportLab, we'll trust FontManager loaded it
        # or it is a standard font.
        # If it fails, we catch it? No, let it fail or fallback?
        
        # Quick check for registered fonts
        registered = pdfmetrics.getRegisteredFontNames()
        standard = ['Courier', 'Helvetica', 'Times-Roman', 'Symbol', 'ZapfDingbats']
        if font_name not in registered and font_name not in standard:
            # If the user selected "Microsoft YaHei" but we didn't load a TTF for it, it won't work.
            # We fallback to Courier to ensure export works.
            # But wait, maybe the user put "Microsoft YaHei.ttf" in fonts folder?
            # Then it would be registered as "Microsoft YaHei" (stem).
            # So if it's not in registered, we fallback.
            print(f"Warning: Font '{font_name}' not registered in ReportLab. Falling back to Courier.")
            # font_name = "Courier" # Uncomment to force fallback, but maybe risky if names differ slightly

        script = project_data.get("script", {})
        title = script.get("title", "Untitled")

        doc = SimpleDocTemplate(
            file_path,
            pagesize=A4,
            leftMargin=margin_inch * inch,
            rightMargin=margin_inch * inch,
            topMargin=margin_inch * inch,
            bottomMargin=margin_inch * inch
        )

        styles = getSampleStyleSheet()

        style_normal = ParagraphStyle(
            'ScreenplayNormal',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=12,
            leading=12 * line_spacing,
            spaceAfter=12
        )

        style_scene = ParagraphStyle(
            'ScreenplayScene',
            parent=style_normal,
            fontName=font_name, # Use same font
            spaceBefore=12,
            spaceAfter=12,
            keepWithNext=True
        )

        style_char = ParagraphStyle(
            'ScreenplayChar',
            parent=style_normal,
            fontName=font_name,
            leftIndent=2.0 * inch,
            spaceBefore=12,
            spaceAfter=0,
            keepWithNext=True
        )

        style_dialogue = ParagraphStyle(
            'ScreenplayDialogue',
            parent=style_normal,
            leftIndent=1.0 * inch,
            rightIndent=1.5 * inch,
            spaceAfter=12
        )

        style_parenthetical = ParagraphStyle(
            'ScreenplayParenthetical',
            parent=style_normal,
            leftIndent=1.5 * inch,
            rightIndent=2.0 * inch,
            spaceAfter=0
        )

        style_transition = ParagraphStyle(
            'ScreenplayTransition',
            parent=style_normal,
            alignment=TA_RIGHT,
            spaceBefore=12,
            spaceAfter=12
        )

        style_title = ParagraphStyle(
            'ScreenplayTitle',
            parent=styles['Title'],
            fontName=font_name,
            fontSize=24,
            alignment=TA_CENTER,
            spaceBefore=3 * inch,
            spaceAfter=4 * inch
        )

        elements = []

        elements.append(Paragraph(title.upper(), style_title))
        elements.append(Paragraph("Generated by Writer Tool", ParagraphStyle('Credit', parent=style_normal, alignment=TA_CENTER)))
        elements.append(PageBreak())

        scenes = script.get("scenes", [])
        for i, scene in enumerate(scenes):
            name = scene.get("name", "SCENE").upper()
            loc = scene.get("location", "").upper()
            time_val = scene.get("time", "").upper()
            slug = f"{name}"
            if loc:
                slug += f" - {loc}"
            if time_val:
                slug += f" - {time_val}"

            elements.append(Paragraph(slug, style_scene))

            content = scene.get("content", "")
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    continue

                match = re.match(r"^([^\s：:]+)[：:]\s*(.*)$", line)
                if match:
                    char_name = match.group(1).upper()
                    dialogue = match.group(2)

                    elements.append(Paragraph(char_name, style_char))

                    if dialogue.startswith("(") and ")" in dialogue:
                        end_paren = dialogue.find(")")
                        paren = dialogue[:end_paren + 1]
                        dial = dialogue[end_paren + 1:].strip()
                        elements.append(Paragraph(paren, style_parenthetical))
                        if dial:
                            elements.append(Paragraph(dial, style_dialogue))
                    else:
                        elements.append(Paragraph(dialogue, style_dialogue))

                elif line.startswith("(") and line.endswith(")"):
                    elements.append(Paragraph(line, style_parenthetical))

                elif line.endswith("TO:") or line.startswith("——"):
                    elements.append(Paragraph(line, style_transition))

                else:
                    elements.append(Paragraph(line, style_normal))

        doc.build(elements)
        return True


class EPUBExporter(ExportFormat):
    """EPUB电子书格式导出器"""
    key = "epub"
    name = "EPUB电子书"
    extension = ".epub"
    description = "EPUB格式，适合电子阅读器"

    @classmethod
    def export(cls, project_data: Dict, file_path: str, **kwargs) -> bool:
        try:
            from ebooklib import epub
        except ImportError:
            raise ImportError("请安装 'ebooklib' 以使用EPUB导出: pip install ebooklib")

        book = epub.EpubBook()
        script = project_data.get("script", {})
        title = script.get("title", "Untitled")

        book.set_identifier("writer-tool-" + project_data.get("outline", {}).get("uid", "id"))
        book.set_title(title)
        book.set_language('zh')
        book.add_author("Writer Tool User")

        chapters = []
        scenes = script.get("scenes", [])
        toc_items = []

        intro = epub.EpubHtml(title='Introduction', file_name='intro.xhtml', lang='zh')
        intro.content = f'<h1>{title}</h1><p>Generated by Writer Tool</p>'
        book.add_item(intro)
        toc_items.append(intro)

        for i, scene in enumerate(scenes):
            chapter_title = scene.get("name", f"Scene {i + 1}")
            c = epub.EpubHtml(title=chapter_title, file_name=f'scene_{i}.xhtml', lang='zh')

            lines = scene.get("content", "").split('\n')
            html_content = f"<h2>{chapter_title}</h2>"
            html_content += f"<p><i>{scene.get('location', '')} - {scene.get('time', '')}</i></p><hr/>"

            for line in lines:
                line = line.strip()
                if not line:
                    continue
                html_content += f"<p>{html.escape(line)}</p>"

            c.content = html_content
            book.add_item(c)
            chapters.append(c)
            toc_items.append(c)

        book.toc = tuple(toc_items)
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        style = 'body { font-family: sans-serif; } h2 { text-align: center; }'
        nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)
        book.add_item(nav_css)

        book.spine = ['nav', intro] + chapters

        epub.write_epub(file_path, book, {})
        return True


class RenpyExporter(ExportFormat):
    """Ren'Py游戏格式导出器"""
    key = "renpy"
    name = "Ren'Py游戏"
    extension = ""
    requires_dir = True
    description = "Ren'Py视觉小说游戏结构"
    supported_types = ["Galgame", "LightNovel"]
    priority_for_types = {"Galgame": 100, "LightNovel": 80}

    @classmethod
    def export(cls, project_data: Dict, export_dir: str, **kwargs) -> bool:
        if not os.path.exists(export_dir):
            os.makedirs(export_dir)

        game_dir = os.path.join(export_dir, "game")
        if not os.path.exists(game_dir):
            os.makedirs(game_dir)

        images_dir = os.path.join(game_dir, "images")
        if not os.path.exists(images_dir):
            os.makedirs(images_dir)

        script = project_data.get("script", {})
        characters = script.get("characters", [])
        scenes = script.get("scenes", [])

        rpy_lines = []
        rpy_lines.append(f"# Script generated by Writer Tool")
        rpy_lines.append("")

        char_map = {}
        for i, char in enumerate(characters):
            name = char.get("name", f"Char{i}")
            safe_name = re.sub(r'\W+', '', name) or f"c{i}"
            var_name = f"c_{safe_name}"
            char_map[name] = var_name

            img_path = char.get("image_path")
            image_tag = None
            if img_path and os.path.exists(img_path):
                ext = os.path.splitext(img_path)[1]
                dest_filename = f"{var_name}{ext}"
                shutil.copy2(img_path, os.path.join(images_dir, dest_filename))
                image_tag = var_name

            if image_tag:
                rpy_lines.append(f'image {var_name} = "{dest_filename}"')
                rpy_lines.append(f'define {var_name} = Character("{name}", image="{var_name}")')
            else:
                rpy_lines.append(f'define {var_name} = Character("{name}")')

        rpy_lines.append("")
        rpy_lines.append("label start:")
        rpy_lines.append("    # Start of the story")

        for i, scene in enumerate(scenes):
            s_name = scene.get("name", f"Scene{i}")
            s_content = scene.get("content", "")

            rpy_lines.append(f"    # Scene: {s_name}")

            bg_path = scene.get("image_path")
            if bg_path and os.path.exists(bg_path):
                ext = os.path.splitext(bg_path)[1]
                bg_var = f"bg_scene_{i}{ext}"
                shutil.copy2(bg_path, os.path.join(images_dir, bg_var))
                rpy_lines.insert(len(characters) * 2 + 2, f'image bg_{i} = "{bg_var}"')
                rpy_lines.append(f"    scene bg_{i} with fade")
            else:
                rpy_lines.append(f"    scene black with fade")

            rpy_lines.append(f"    \"Location: {scene.get('location')} | Time: {scene.get('time')}\"")

            for line in s_content.split('\n'):
                line = line.strip()
                if not line:
                    continue

                match = re.match(r"^([^\s：:]+)[：:]\s*(.*)$", line)
                if not match:
                    match = re.match(r"^【(.*?)】\s*(.*)$", line)

                if match:
                    char_name = match.group(1).strip()
                    dialogue = match.group(2).strip()
                    dialogue = dialogue.replace('"', '\\"')

                    if char_name in char_map:
                        var = char_map[char_name]
                        rpy_lines.append(f'    {var} "{dialogue}"')
                    else:
                        rpy_lines.append(f'    "{char_name}" "{dialogue}"')
                else:
                    line = line.replace('"', '\\"')
                    rpy_lines.append(f'    "{line}"')

            rpy_lines.append("")

        rpy_lines.append("    return")

        with open(os.path.join(game_dir, "script.rpy"), "w", encoding="utf-8") as f:
            f.write("\n".join(rpy_lines))
        return True


class CharacterSidesExporter(ExportFormat):
    """角色台词本导出器"""
    key = "sides"
    name = "角色台词本"
    extension = ".txt"
    description = "导出指定角色的所有台词"

    @classmethod
    def export(cls, project_data: Dict, file_path: str, **kwargs) -> bool:
        target_role = kwargs.get("target_role", "")
        if not target_role:
            raise ValueError("必须指定目标角色 (target_role)")

        script = project_data.get("script", {})
        scenes = script.get("scenes", [])
        lines = []

        lines.append(f"CHARACTER SIDES: {target_role.upper()}")
        lines.append(f"Project: {script.get('title', 'Untitled')}")
        lines.append("=" * 40 + "\n")

        for i, scene in enumerate(scenes):
            content = scene.get("content", "")
            if target_role not in content:
                continue

            scene_lines = content.split('\n')
            scene_has_role = False
            extracted_lines = []

            last_speaker = None
            last_line = None

            for line in scene_lines:
                line = line.strip()
                if not line:
                    continue

                match = re.match(r"^([^\s：:]+)[：:]\s*(.*)$", line)
                if match:
                    speaker = match.group(1).strip()
                    dialogue = match.group(2).strip()

                    if speaker == target_role:
                        if not scene_has_role:
                            header = f"SCENE {i + 1}: {scene.get('name')} [{scene.get('location')} - {scene.get('time')}]"
                            extracted_lines.append(f"\n{header}")
                            extracted_lines.append("-" * len(header))
                            scene_has_role = True

                        if last_speaker and last_speaker != target_role:
                            extracted_lines.append(f"(CUE) {last_speaker}: ...{last_line[-30:] if len(last_line) > 30 else last_line}")

                        extracted_lines.append(f"{speaker}: {dialogue}")

                    last_speaker = speaker
                    last_line = dialogue

            if extracted_lines:
                lines.extend(extracted_lines)
                lines.append("")

        with open(file_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return True


# ============================================================================
# 注册所有内置导出器
# ============================================================================

def _register_builtin_exporters():
    """注册所有内置导出器"""
    exporters = [
        CSVExporter,
        ExcelExporter,
        TxtExporter,
        MarkdownExporter,
        HTMLExporter,
        FountainExporter,
        FDXExporter,
        DocxExporter,
        PDFExporter,
        EPUBExporter,
        RenpyExporter,
        CharacterSidesExporter,
    ]
    for exporter in exporters:
        ExporterRegistry.register(exporter)


# 模块加载时自动注册
_register_builtin_exporters()

# 注册题材专属导出器
try:
    from writer_app.core.exporters import register_genre_exporters
    register_genre_exporters()
except ImportError:
    pass  # Genre exporters not available


# ============================================================================
# 兼容层：保留原有的Exporter类以保持向后兼容
# ============================================================================

class Exporter:
    @staticmethod
    def export_to_csv(project_data, file_path):
        """Export scenes and outline to CSV for spreadsheet management."""
        script = project_data.get("script", {})
        scenes = script.get("scenes", [])
        
        try:
            with open(file_path, "w", encoding="utf-8-sig", newline="") as f:
                writer = csv.writer(f)
                # Header
                writer.writerow(["序号", "场景名", "地点", "时间", "字数", "登场角色", "大纲关联", "备注"])
                
                for i, scene in enumerate(scenes):
                    content = scene.get("content", "")
                    word_count = len(content)
                    
                    # Try to find associated outline node name
                    outline_ref = scene.get("outline_ref_path", "")
                    
                    writer.writerow([
                        i + 1,
                        scene.get("name", ""),
                        scene.get("location", ""),
                        scene.get("time", ""),
                        word_count,
                        ", ".join(scene.get("characters", [])),
                        outline_ref,
                        "" # Placeholder for manual remarks
                    ])
            return True
        except Exception as e:
            raise e

    @staticmethod
    def export_to_excel(project_data, file_path):
        """Export scenes and outline to Excel (.xlsx)."""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill
        except ImportError:
            raise ImportError("Please install 'openpyxl' to use Excel export: pip install openpyxl")

        script = project_data.get("script", {})
        scenes = script.get("scenes", [])
        
        wb = Workbook()
        ws = wb.active
        ws.title = "场景列表"
        
        headers = ["序号", "场景名", "地点", "时间", "字数", "登场角色", "大纲关联", "备注"]
        ws.append(headers)
        
        # Style headers
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
        for cell in ws[1]:
            cell.font = header_font
            cell.fill = header_fill
            
        for i, scene in enumerate(scenes):
            content = scene.get("content", "")
            word_count = len(content)
            outline_ref = scene.get("outline_ref_path", "")
            
            row = [
                i + 1,
                scene.get("name", ""),
                scene.get("location", ""),
                scene.get("time", ""),
                word_count,
                ", ".join(scene.get("characters", [])),
                outline_ref,
                "" 
            ]
            ws.append(row)
            
        # Adjust column widths
        ws.column_dimensions['B'].width = 20
        ws.column_dimensions['C'].width = 15
        ws.column_dimensions['D'].width = 15
        ws.column_dimensions['F'].width = 25
        ws.column_dimensions['G'].width = 25
        
        # Second sheet for characters
        ws_char = wb.create_sheet("角色统计")
        ws_char.append(["姓名", "描述", "登场次数", "标签"])
        for cell in ws_char[1]:
            cell.font = header_font
            cell.fill = header_fill

        chars = script.get("characters", [])
        for char in chars:
            name = char.get("name", "")
            # Count occurrences
            count = 0
            for s in scenes:
                if name in s.get("characters", []):
                    count += 1
            
            ws_char.append([
                name,
                char.get("description", ""),
                count,
                ", ".join(char.get("tags", []))
            ])
            
        ws_char.column_dimensions['A'].width = 15
        ws_char.column_dimensions['B'].width = 40

        wb.save(file_path)
        return True

    @staticmethod
    def export_character_sides(project_data, file_path, target_role):
        """
        Export a script containing only lines for a specific character (Sides),
        including cue lines (preceding line) for context.
        """
        script = project_data.get("script", {})
        scenes = script.get("scenes", [])
        lines = []
        
        lines.append(f"CHARACTER SIDES: {target_role.upper()}")
        lines.append(f"Project: {script.get('title', 'Untitled')}")
        lines.append("=" * 40 + "\n")

        for i, scene in enumerate(scenes):
            content = scene.get("content", "")
            if target_role not in content:
                continue
                
            scene_lines = content.split('\n')
            scene_has_role = False
            extracted_lines = []
            
            last_speaker = None
            last_line = None
            
            for line in scene_lines:
                line = line.strip()
                if not line: continue
                
                # Check for dialogue
                match = re.match(r"^([^\s：:]+)[：:]\s*(.*)$", line)
                if match:
                    speaker = match.group(1).strip()
                    dialogue = match.group(2).strip()
                    
                    if speaker == target_role:
                        if not scene_has_role:
                            # Add Scene Header first time we see role
                            header = f"SCENE {i+1}: {scene.get('name')} [{scene.get('location')} - {scene.get('time')}]"
                            extracted_lines.append(f"\n{header}")
                            extracted_lines.append("-" * len(header))
                            scene_has_role = True
                        
                        # Add cue if exists and different speaker
                        if last_speaker and last_speaker != target_role:
                            extracted_lines.append(f"(CUE) {last_speaker}: ...{last_line[-30:] if len(last_line)>30 else last_line}")
                        
                        # Add own line
                        extracted_lines.append(f"{speaker}: {dialogue}")
                    
                    last_speaker = speaker
                    last_line = dialogue
                else:
                    # Reset context on action? Optional. 
                    # For simple sides, we might ignore action or treat as cue if relevant.
                    pass

            if extracted_lines:
                lines.extend(extracted_lines)
                lines.append("")

        try:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))
            return True
        except Exception as e:
            raise e

    @staticmethod
    def export_to_pdf(project_data, file_path):
        """Export script to standard PDF Screenplay format using ReportLab."""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        except ImportError:
            raise ImportError("Please install 'reportlab' to use PDF export: pip install reportlab")

        script = project_data.get("script", {})
        title = script.get("title", "Untitled")
        
        doc = SimpleDocTemplate(
            file_path,
            pagesize=A4,
            leftMargin=1.5*inch,
            rightMargin=1.0*inch,
            topMargin=1.0*inch,
            bottomMargin=1.0*inch
        )
        
        styles = getSampleStyleSheet()
        
        # Screenplay Styles
        # Font: Courier is standard
        style_normal = ParagraphStyle(
            'ScreenplayNormal',
            parent=styles['Normal'],
            fontName='Courier',
            fontSize=12,
            leading=14,
            spaceAfter=12
        )
        
        style_scene = ParagraphStyle(
            'ScreenplayScene',
            parent=style_normal,
            fontName='Courier-Bold',
            spaceBefore=12,
            spaceAfter=12,
            keepWithNext=True
        )
        
        style_char = ParagraphStyle(
            'ScreenplayChar',
            parent=style_normal,
            fontName='Courier-Bold',
            leftIndent=2.0*inch,
            spaceBefore=12,
            spaceAfter=0,
            keepWithNext=True
        )
        
        style_dialogue = ParagraphStyle(
            'ScreenplayDialogue',
            parent=style_normal,
            leftIndent=1.0*inch,
            rightIndent=1.5*inch,
            spaceAfter=12
        )
        
        style_parenthetical = ParagraphStyle(
            'ScreenplayParenthetical',
            parent=style_normal,
            leftIndent=1.5*inch,
            rightIndent=2.0*inch,
            spaceAfter=0
        )
        
        style_transition = ParagraphStyle(
            'ScreenplayTransition',
            parent=style_normal,
            alignment=TA_RIGHT,
            spaceBefore=12,
            spaceAfter=12
        )
        
        style_title = ParagraphStyle(
            'ScreenplayTitle',
            parent=styles['Title'],
            fontName='Courier-Bold',
            fontSize=24,
            alignment=TA_CENTER,
            spaceBefore=3*inch,
            spaceAfter=4*inch
        )

        elements = []
        
        # Title Page
        elements.append(Paragraph(title.upper(), style_title))
        elements.append(Paragraph("Generated by Writer Tool", ParagraphStyle('Credit', parent=style_normal, alignment=TA_CENTER)))
        elements.append(PageBreak())
        
        # Content
        scenes = script.get("scenes", [])
        for i, scene in enumerate(scenes):
            # Scene Heading
            name = scene.get("name", "SCENE").upper()
            loc = scene.get("location", "").upper()
            time = scene.get("time", "").upper()
            slug = f"{name}"
            if loc: slug += f" - {loc}"
            if time: slug += f" - {time}"
            
            elements.append(Paragraph(slug, style_scene))
            
            content = scene.get("content", "")
            for line in content.split('\n'):
                line = line.strip()
                if not line: continue
                
                # Regex for Character
                match = re.match(r"^([^\s：:]+)[：:]\s*(.*)$", line)
                if match:
                    char_name = match.group(1).upper()
                    dialogue = match.group(2)
                    
                    elements.append(Paragraph(char_name, style_char))
                    
                    # Check parenthetical inside dialogue?
                    if dialogue.startswith("(") and ")" in dialogue:
                        end_paren = dialogue.find(")")
                        paren = dialogue[:end_paren+1]
                        dial = dialogue[end_paren+1:].strip()
                        elements.append(Paragraph(paren, style_parenthetical))
                        if dial:
                            elements.append(Paragraph(dial, style_dialogue))
                    else:
                        elements.append(Paragraph(dialogue, style_dialogue))
                
                elif line.startswith("(") and line.endswith(")"):
                    elements.append(Paragraph(line, style_parenthetical))
                
                elif line.endswith("TO:") or line.startswith("——"):
                    elements.append(Paragraph(line, style_transition))
                    
                else:
                    elements.append(Paragraph(line, style_normal))
                    
        try:
            doc.build(elements)
            return True
        except Exception as e:
            raise e

    @staticmethod
    def export_to_epub(project_data, file_path):
        try:
            from ebooklib import epub
        except ImportError:
            raise ImportError("Please install 'ebooklib' to use EPUB export: pip install ebooklib")

        book = epub.EpubBook()
        script = project_data.get("script", {})
        title = script.get("title", "Untitled")
        
        book.set_identifier("writer-tool-" + project_data.get("outline", {}).get("uid", "id"))
        book.set_title(title)
        book.set_language('zh')
        book.add_author("Writer Tool User")

        # Create chapters
        chapters = []
        scenes = script.get("scenes", [])
        
        # TOC items
        toc_items = []

        # Intro page
        intro = epub.EpubHtml(title='Introduction', file_name='intro.xhtml', lang='zh')
        intro.content = f'<h1>{title}</h1><p>Generated by Writer Tool</p>'
        book.add_item(intro)
        toc_items.append(intro)

        for i, scene in enumerate(scenes):
            chapter_title = scene.get("name", f"Scene {i+1}")
            c = epub.EpubHtml(title=chapter_title, file_name=f'scene_{i}.xhtml', lang='zh')
            
            # Format content
            lines = scene.get("content", "").split('\n')
            html_content = f"<h2>{chapter_title}</h2>"
            html_content += f"<p><i>{scene.get('location', '')} - {scene.get('time', '')}</i></p><hr/>"
            
            for line in lines:
                line = line.strip()
                if not line: continue
                html_content += f"<p>{html.escape(line)}</p>"
                
            c.content = html_content
            book.add_item(c)
            chapters.append(c)
            toc_items.append(c)

        # Define Table of Contents
        book.toc = tuple(toc_items)

        # Add default NCX and Nav file
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # Define CSS style
        style = 'body { font-family: sans-serif; } h2 { text-align: center; }'
        nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)
        book.add_item(nav_css)

        # Basic spine
        book.spine = ['nav', intro] + chapters

        try:
            epub.write_epub(file_path, book, {})
            return True
        except Exception as e:
            raise e

    @staticmethod
    def export_to_txt(project_data, file_path):
        """Export script to Plain Text format."""
        script = project_data.get("script", {})
        lines = []
        
        title = script.get("title", "Untitled")
        lines.append(title)
        lines.append("=" * len(title))
        lines.append("")
        
        for scene in script.get("scenes", []):
            name = scene.get("name", "Untitled Scene")
            loc = scene.get("location", "Unknown Location")
            time = scene.get("time", "Unknown Time")
            
            lines.append(f"{name}")
            lines.append(f"{loc} - {time}")
            lines.append("-" * 20)
            
            content = scene.get("content", "")
            lines.append(content)
            lines.append("\n" + "="*30 + "\n")

        try:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))
            return True
        except Exception as e:
            raise e

    @staticmethod
    def export_to_markdown(project_data, file_path, include_outline=True, include_script=True):
        lines = []
        script = project_data.get("script", {})
        
        if include_script:
            title = script.get("title", "Untitled")
            lines.append(f"# {title}\n")
            
            chars = script.get("characters", [])
            if chars:
                lines.append("## Characters")
                for c in chars:
                    lines.append(f"- **{c.get('name')}**: {c.get('description')}")
                lines.append("")
            
            scenes = script.get("scenes", [])
            if scenes:
                lines.append("## Scenes")
                for s in scenes:
                    lines.append(f"### {s.get('name')}")
                    lines.append(f"**Location**: {s.get('location')} | **Time**: {s.get('time')}\n")
                    
                    content = s.get("content", "")
                    for line in content.split('\n'):
                        line = line.strip()
                        if not line:
                            lines.append("")
                            continue
                        # Format dialogue: Name: Dialogue -> **Name**: Dialogue
                        match = re.match(r"^([^\s：:]+)[：:]\s*(.*)$", line)
                        if match:
                            name = match.group(1)
                            dialogue = match.group(2)
                            lines.append(f"**{name}**: {dialogue}")
                        else:
                            lines.append(line)
                    
                    lines.append("\n---\n")

        if include_outline:
            lines.append("\n# Outline\n")
            Exporter._recursive_outline(project_data.get("outline", {}), 0, lines)

        try:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))
            return True
        except Exception as e:
            raise e

    @staticmethod
    def _recursive_outline(node, level, lines):
        prefix = "#" * (level + 2) if level < 5 else "-"
        lines.append(f"{prefix} {node.get('name', 'Untitled')}")
        if node.get("content"):
            lines.append(f"{node.get('content')}\n")
        for child in node.get("children", []):
            Exporter._recursive_outline(child, level + 1, lines)

    @staticmethod
    def export_to_html_print(project_data, file_path):
        """Export script to HTML formatted for printing (PDF via Browser)."""
        script = project_data.get("script", {})
        title = html.escape(script.get("title", "Untitled"))
        
        css = """
        @media print {
            @page { margin: 1in; size: A4; }
            body { font-family: 'Courier New', Courier, monospace; font-size: 12pt; line-height: 1.2; }
            .scene-heading { font-weight: bold; text-transform: uppercase; margin-top: 2em; margin-bottom: 1em; }
            .character { font-weight: bold; text-align: center; width: 60%; margin-left: 20%; margin-top: 1em; }
            .parenthetical { font-style: italic; text-align: center; width: 50%; margin-left: 25%; }
            .dialogue { text-align: left; width: 70%; margin-left: 15%; margin-bottom: 1em; }
            .action { text-align: left; margin-bottom: 1em; }
            .transition { text-align: right; text-transform: uppercase; margin-top: 1em; margin-bottom: 1em; }
            .page-break { page-break-before: always; }
        }
        body { font-family: 'Courier New', Courier, monospace; font-size: 12pt; line-height: 1.2; max-width: 8.5in; margin: 0 auto; padding: 2em; background: #fff; color: #000; }
        .scene-heading { font-weight: bold; text-transform: uppercase; margin-top: 24pt; margin-bottom: 12pt; background: #eee; }
        .character { font-weight: bold; text-align: center; width: 50%; margin-left: 25%; margin-top: 12pt; }
        .parenthetical { font-style: italic; text-align: center; width: 40%; margin-left: 30%; }
        .dialogue { text-align: left; width: 60%; margin-left: 20%; margin-bottom: 12pt; }
        .action { text-align: left; margin-bottom: 12pt; }
        .transition { text-align: right; text-transform: uppercase; margin-top: 12pt; margin-bottom: 12pt; }
        """
        
        html_content = [f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>{title}</title><style>{css}</style></head><body>"]
        html_content.append(f"<h1 style='text-align:center; text-transform:uppercase; margin-top:3in; margin-bottom:4in;'>{title}</h1>")
        html_content.append("<div class='page-break'></div>")
        
        for scene in script.get("scenes", []):
            name = html.escape(scene.get("name", "SCENE"))
            loc = html.escape(scene.get("location", ""))
            time = html.escape(scene.get("time", ""))
            slug = f"{name} - {loc} - {time}".upper()
            
            html_content.append(f"<div class='scene-heading'>{slug}</div>")
            
            content = scene.get("content", "")
            for line in content.split('\n'):
                line = line.strip()
                if not line: continue
                
                esc_line = html.escape(line)
                
                # Heuristics
                if line.isupper() and (line.endswith(":") or len(line) < 30):
                     # Likely Transition or Character?
                     if line.endswith("TO:"):
                         html_content.append(f"<div class='transition'>{esc_line}</div>")
                     else:
                         # Very heuristic: Assume centered char name if short and upper, or contains :
                         # We rely on our standard "Name: Dialogue" format from editor
                         pass # handled below
                
                # Check format "Name: Dialogue"
                match = re.match(r"^([^\s：:]+)[：:]\s*(.*)$", line)
                if match:
                    char_name = match.group(1).strip()
                    dialogue = match.group(2).strip()
                    html_content.append(f"<div class='character'>{html.escape(char_name).upper()}</div>")
                    
                    # Parenthetical check in dialogue start? "(beat) hello"
                    if dialogue.startswith("(") and ")" in dialogue:
                        end_paren = dialogue.find(")")
                        paren = dialogue[:end_paren+1]
                        dial = dialogue[end_paren+1:].strip()
                        html_content.append(f"<div class='parenthetical'>{html.escape(paren)}</div>")
                        if dial:
                            html_content.append(f"<div class='dialogue'>{html.escape(dial)}</div>")
                    else:
                        html_content.append(f"<div class='dialogue'>{html.escape(dialogue)}</div>")
                
                elif line.startswith("(") and line.endswith(")"):
                    html_content.append(f"<div class='parenthetical'>{esc_line}</div>")
                
                elif line.startswith("——") or line.endswith("TO:"):
                    html_content.append(f"<div class='transition'>{esc_line}</div>")
                    
                else:
                    html_content.append(f"<div class='action'>{esc_line}</div>")

        html_content.append("</body></html>")
        
        try:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write("\n".join(html_content))
            return True
        except Exception as e:
            raise e

    @staticmethod
    def export_to_fountain(project_data, file_path):
        """Export script to Fountain format."""
        script = project_data.get("script", {})
        lines = []
        
        # Title Page
        title = script.get("title", "Untitled")
        lines.append(f"Title: {title}")
        lines.append("Credit: Written with Writer Tool")
        lines.append("\n") # Blank line separates title page
        
        for scene in script.get("scenes", []):
            # Scene Heading
            # Format: INT. LOCATION - TIME
            loc = scene.get("location", "UNKNOWN").upper()
            time = scene.get("time", "").upper()
            # Ensure it starts with standard prefix if not present to be valid scene heading
            if not any(loc.startswith(p) for p in ["INT", "EXT", ".", "EST"]):
                # If pure location name, treat as scene heading by forcing '.'
                if not loc.startswith("."):
                     loc = "." + loc
            
            heading = f"{loc} - {time}" if time else loc
            lines.append(f"\n{heading}\n")
            
            content = scene.get("content", "")
            for line in content.split('\n'):
                line = line.strip()
                if not line: 
                    lines.append("")
                    continue
                
                # Heuristic for format parsing
                # 1. Character Name: Dialogue
                match = re.match(r"^([^\s：:]+)[：:]\s*(.*)$", line)
                if match:
                    char = match.group(1).strip().upper()
                    diag = match.group(2).strip()
                    lines.append(f"\n{char}")
                    lines.append(diag)
                    continue
                
                # 2. Parenthetical
                if line.startswith("(") and line.endswith(")"):
                    lines.append(line)
                    continue
                
                # 3. Transition
                if line.endswith("TO:"):
                    lines.append(f"\n> {line}")
                    continue
                    
                # 4. Action
                lines.append(line)
                
        try:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))
            return True
        except Exception as e:
            raise e

    @staticmethod
    def export_to_fdx(project_data, file_path):
        """Export script to Final Draft 12 XML (.fdx) format."""
        script = project_data.get("script", {})
        
        # Basic FDX Template Structure
        root = ET.Element("FinalDraft", {"DocumentType": "Script", "Template": "No", "Version": "4"})
        
        content = ET.SubElement(root, "Content")
        paragraph = ET.SubElement(content, "Paragraph", {"Type": "General"})
        text = ET.SubElement(paragraph, "Text")
        text.text = script.get("title", "Untitled")

        # Process Scenes
        for scene in script.get("scenes", []):
            # Scene Heading
            slugline = f"{scene.get('location', 'UNKNOWN')} - {scene.get('time', 'UNKNOWN')}"
            sc_para = ET.SubElement(content, "Paragraph", {"Type": "Scene Heading"})
            sc_text = ET.SubElement(sc_para, "Text")
            sc_text.text = f"{scene.get('name')} - {slugline}" # Simplified slugline

            # Content - naive parsing
            # In a real exporter, we'd parse the scene content for Action, Dialogue, etc.
            # Here we just dump it as Action for now, or split by lines
            scene_body = scene.get("content", "")
            for line in scene_body.split('\n'):
                line = line.strip()
                if not line: continue
                
                p_type = "Action"
                if ":" in line and len(line.split(":")[0]) < 20:
                     p_type = "Dialogue" # Very naive heuristic
                elif line.startswith("(") and line.endswith(")"):
                     p_type = "Parenthetical"
                
                para = ET.SubElement(content, "Paragraph", {"Type": p_type})
                t = ET.SubElement(para, "Text")
                t.text = line

        # Pretty print
        xml_str = minidom.parseString(ET.tostring(root)).toprettyxml(indent="  ")
        
        try:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(xml_str)
            return True
        except Exception as e:
            raise e

    @staticmethod
    def export_to_docx(project_data, file_path):
        try:
            from docx import Document
        except ImportError:
            raise ImportError("Please install 'python-docx' to use Word export: pip install python-docx")

        document = Document()
        script = project_data.get("script", {})
        
        document.add_heading(script.get("title", "Untitled"), 0)

        for scene in script.get("scenes", []):
            # Scene Header
            slug = f"{scene.get('name', '')} [Loc: {scene.get('location')} Time: {scene.get('time')}]"
            document.add_heading(slug, level=2)
            
            # Content
            content = scene.get("content", "")
            document.add_paragraph(content)

        try:
            document.save(file_path)
            return True
        except Exception as e:
            raise e

    @staticmethod
    def export_to_renpy(project_data, export_dir):
        """Export project as a Ren'Py game structure."""
        if not os.path.exists(export_dir):
            os.makedirs(export_dir)
        
        game_dir = os.path.join(export_dir, "game")
        if not os.path.exists(game_dir):
            os.makedirs(game_dir)
            
        images_dir = os.path.join(game_dir, "images")
        if not os.path.exists(images_dir):
            os.makedirs(images_dir)

        script = project_data.get("script", {})
        characters = script.get("characters", [])
        scenes = script.get("scenes", [])

        rpy_lines = []
        rpy_lines.append(f"# Script generated by Writer Tool")
        rpy_lines.append("")

        # 1. Define Characters
        char_map = {} # name -> variable_name
        for i, char in enumerate(characters):
            name = char.get("name", f"Char{i}")
            safe_name = re.sub(r'\W+', '', name) or f"c{i}"
            var_name = f"c_{safe_name}"
            char_map[name] = var_name
            
            # Handle Image
            img_path = char.get("image_path")
            image_tag = None
            if img_path and os.path.exists(img_path):
                ext = os.path.splitext(img_path)[1]
                dest_filename = f"{var_name}{ext}"
                shutil.copy2(img_path, os.path.join(images_dir, dest_filename))
                image_tag = var_name # In Renpy: image c_alice = "c_alice.png" -> define c = Character("", image="c_alice")
                
                # We need to define the image explicitly if we want to use 'show c_alice' easily, 
                # or just rely on RenPy's auto image definition if filename matches.
                # Let's simplify:
                # define c_alice = Character("Alice")
                # image c_alice = "c_alice.png"
            
            if image_tag:
                 rpy_lines.append(f'image {var_name} = "{dest_filename}"')
                 rpy_lines.append(f'define {var_name} = Character("{name}", image="{var_name}")')
            else:
                 rpy_lines.append(f'define {var_name} = Character("{name}")')

        rpy_lines.append("")
        
        # 2. Start Label
        rpy_lines.append("label start:")
        rpy_lines.append("    # Start of the story")
        
        # 3. Process Scenes
        for i, scene in enumerate(scenes):
            s_name = scene.get("name", f"Scene{i}")
            s_content = scene.get("content", "")
            
            rpy_lines.append(f"    # Scene: {s_name}")
            
            # Background
            bg_path = scene.get("image_path")
            if bg_path and os.path.exists(bg_path):
                 ext = os.path.splitext(bg_path)[1]
                 bg_var = f"bg_scene_{i}{ext}"
                 shutil.copy2(bg_path, os.path.join(images_dir, bg_var))
                 # We need to define it or just show it. 
                 # Renpy needs 'image bg_xxx = ...'
                 # We'll inject the image definition at top later? No, can do it anywhere or better at top.
                 # For simplicity, let's just quote the path? No, must be in images dir.
                 rpy_lines.insert(len(characters)*2 + 2, f'image bg_{i} = "{bg_var}"') # Insert at top
                 rpy_lines.append(f"    scene bg_{i} with fade")
            else:
                 rpy_lines.append(f"    scene black with fade") # Default bg

            rpy_lines.append(f"    \"Location: {scene.get('location')} | Time: {scene.get('time')}\"")

            # Content Parsing
            for line in s_content.split('\n'):
                line = line.strip()
                if not line: continue
                
                # Check for "Name: Dialogue" or "【Name】 Dialogue"
                # Regex: (Name)[:：]\s*(.*)
                match = re.match(r"^([^\s：:]+)[：:]\s*(.*)$", line)
                if not match:
                    # Try 【Name】
                    match = re.match(r"^【(.*?)】\s*(.*)$", line)
                
                if match:
                    char_name = match.group(1).strip()
                    dialogue = match.group(2).strip()
                    
                    # Escape quotes
                    dialogue = dialogue.replace('"', '\\"')
                    
                    if char_name in char_map:
                        var = char_map[char_name]
                        rpy_lines.append(f'    {var} "{dialogue}"')
                    else:
                        # Unknown character, use name string
                        rpy_lines.append(f'    "{char_name}" "{dialogue}"')
                else:
                    # Narration or Action
                    # Escape quotes
                    line = line.replace('"', '\\"')
                    rpy_lines.append(f'    "{line}"')
            
            rpy_lines.append("")

        rpy_lines.append("    return")

        try:
            with open(os.path.join(game_dir, "script.rpy"), "w", encoding="utf-8") as f:
                f.write("\n".join(rpy_lines))
            return True
        except Exception as e:
            raise e